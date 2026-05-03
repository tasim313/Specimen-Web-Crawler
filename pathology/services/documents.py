from __future__ import annotations

import logging
import re
from pathlib import Path

from .constants import SECTION_HEADINGS, SIZE_PATTERN
from .normalizers import (
    build_specimen_name,
    clean_whitespace,
    infer_organ_name,
    infer_specimen_type,
    normalize_specimen_size,
)
from .types import ParsedSpecimenData

logger = logging.getLogger("pathology.documents")


def parse_document(file_path: Path, category: str) -> ParsedSpecimenData | None:
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".docx":
            content = _parse_docx_text(file_path)
        elif suffix == ".pdf":
            content = _parse_pdf_text(file_path)
        else:
            logger.warning("Unsupported file type skipped: %s", file_path)
            return None
    except Exception:
        logger.exception("Failed to parse document: %s", file_path)
        return None

    if not content:
        logger.warning("No content extracted from %s", file_path)
        return None

    raw_specimen_name = _extract_specimen_name(content, file_path)
    procedure_block = _extract_section(content, "procedure")
    size_context = "\n".join(
        part
        for part in (
            _extract_section(content, "gross description"),
            _extract_section(content, "specimen"),
            content[:4000],
        )
        if part
    )

    specimen_type = infer_specimen_type(
        raw_specimen_name,
        procedure_block,
        content[:2000],
    )
    organ_name = infer_organ_name(
        raw_specimen_name,
        category.replace("_", " ").title(),
        source_stem=file_path.stem,
    )
    specimen_name = build_specimen_name(
        raw_specimen_name,
        organ_name,
        specimen_type,
        source_stem=file_path.stem,
    )
    site_name = _extract_site_name(content, organ_name)
    laterality = _extract_laterality(content, site_name)
    specimen_size = normalize_specimen_size(
        " ".join(part for part in (raw_specimen_name, procedure_block) if part),
        specimen_type,
        organ_name,
    ) or _extract_size(size_context)

    return ParsedSpecimenData(
        specimen_name=specimen_name,
        organ_name=organ_name,
        site_name=site_name,
        laterality=laterality,
        specimen_type=specimen_type,
        specimen_size=specimen_size,
        source_site="cap.org",
        source_file=file_path,
    )


def _parse_docx_text(file_path: Path) -> str:
    from docx import Document

    document = Document(file_path)
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _parse_pdf_text(file_path: Path) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        if clean_whitespace(text):
            return text
    except Exception:
        logger.exception("PyPDF2 failed for %s; attempting pdfplumber fallback", file_path)

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception:
        logger.exception("pdfplumber fallback failed for %s", file_path)
        raise


def _extract_specimen_name(content: str, file_path: Path) -> str:
    normalized_content = clean_whitespace(content[:5000])
    protocol_match = re.search(
        r"(Protocol for the Examination of .*?)(?:Version:|CAP Laboratory|For accreditation purposes)",
        normalized_content,
        flags=re.IGNORECASE,
    )
    if protocol_match:
        return clean_whitespace(protocol_match.group(1))

    case_summary_match = re.search(
        r"CASE SUMMARY:\s*\((.*?)\)",
        normalized_content,
        flags=re.IGNORECASE,
    )
    if case_summary_match:
        return clean_whitespace(case_summary_match.group(1))

    skip_prefixes = (
        "©",
        "version:",
        "protocol posting date:",
        "cap laboratory",
        "authors",
        "summary of changes",
    )
    for line in content.splitlines():
        cleaned = clean_whitespace(line)
        if len(cleaned) > 5 and not cleaned.lower().startswith(skip_prefixes):
            return cleaned
    return file_path.stem.replace("_", " ").strip()


def _extract_section(content: str, heading: str) -> str:
    lines = [clean_whitespace(line) for line in content.splitlines()]
    normalized_heading = heading.lower()

    for index, line in enumerate(lines):
        if normalized_heading in line.lower():
            window = []
            for candidate in lines[index:index + 12]:
                candidate_lower = candidate.lower()
                if (
                    window
                    and any(
                        section in candidate_lower and section != normalized_heading
                        for section in SECTION_HEADINGS
                    )
                ):
                    break
                window.append(candidate)
            return "\n".join(window)
    return ""


def _extract_size(content: str) -> str:
    match = re.search(SIZE_PATTERN, content, flags=re.IGNORECASE)
    return match.group(0) if match else ""


def _extract_site_name(content: str, organ_name: str) -> str:
    lines = [clean_whitespace(line) for line in content.splitlines()]
    options = _extract_option_values(lines, "tumor site", max_options=16)
    if not options:
        return organ_name
    return "; ".join(options)


def _extract_laterality(content: str, site_name: str) -> str:
    lines = [clean_whitespace(line) for line in content.splitlines()]
    options = _extract_option_values(lines, "specimen laterality", max_options=8)
    if not options:
        laterality_terms = []
        haystack = site_name.lower()
        for label in ("Right", "Left", "Bilateral", "Not specified", "Cannot be determined"):
            if label.lower() in haystack:
                laterality_terms.append(label)
        options = laterality_terms
    return "; ".join(options)


def _extract_option_values(lines: list[str], heading_text: str, max_options: int = 12) -> list[str]:
    for index, line in enumerate(lines):
        if heading_text not in line.lower():
            continue
        options = []
        started = False
        for candidate in lines[index + 1:index + 30]:
            if not candidate:
                continue
            if candidate.startswith("___"):
                started = True
                label = _clean_option_label(candidate)
                if label and label not in options:
                    options.append(label)
                if len(options) >= max_options:
                    break
                continue
            if started:
                break
        if options:
            return options
    return []


def _clean_option_label(value: str) -> str:
    cleaned = re.sub(r"^_+\s*", "", value).strip()
    cleaned = re.sub(r":\s*_+$", "", cleaned)
    cleaned = re.sub(r":\s*[_\s]+$", "", cleaned)
    cleaned = cleaned.rstrip(":").strip()
    return cleaned
